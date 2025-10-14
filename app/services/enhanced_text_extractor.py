"""
Enhanced Text Extraction System for Resume Screener
Handles PDF, DOCX, and TXT files with improved text quality and formatting preservation.
"""

import os
import re
import tempfile
import io
from typing import Dict, List, Optional, Tuple
import logging

# Multiple PDF extraction libraries for fallback
import pdfplumber
import fitz  # PyMuPDF
import PyPDF2
import docx2txt
from docx import Document

# Setup logging
logger = logging.getLogger(__name__)

class EnhancedTextExtractor:
    """
    Advanced text extractor that handles various document formats with improved quality.
    Uses multiple extraction methods and post-processing to handle text bisection issues.
    """
    
    def __init__(self):
        self.extraction_methods = {
            '.pdf': [self._extract_pdf_pdfplumber, self._extract_pdf_pymupdf, self._extract_pdf_pypdf2],
            '.docx': [self._extract_docx_python_docx, self._extract_docx_docx2txt],
            '.txt': [self._extract_txt]
        }
    
    def extract_text(self, file_path: str) -> Dict[str, str]:
        """
        Extract text from file with enhanced processing and error recovery.
        
        Args:
            file_path: Path to the file to extract text from
            
        Returns:
            Dictionary containing:
            - 'raw_text': Cleaned raw text
            - 'formatted_text': Text with preserved formatting
            - 'extraction_method': Method used for extraction
            - 'quality_score': Text quality assessment (0-1)
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext not in self.extraction_methods:
            raise ValueError(f"Unsupported file format: {ext}")
        
        best_result = None
        best_score = 0
        
        # Try multiple extraction methods and pick the best
        for method in self.extraction_methods[ext]:
            try:
                result = method(file_path)
                quality_score = self._assess_text_quality(result['raw_text'])
                result['quality_score'] = quality_score
                
                if quality_score > best_score:
                    best_result = result
                    best_score = quality_score
                    
                # If we get high quality, don't try other methods
                if quality_score > 0.85:
                    break
                    
            except Exception as e:
                logger.warning(f"Extraction method {method.__name__} failed: {str(e)}")
                continue
        
        if not best_result:
            raise Exception("All extraction methods failed")
        
        return best_result
    
    def extract_text_from_bytes(self, file_bytes: bytes, file_extension: str) -> Dict[str, str]:
        """
        Extract text from file bytes (for blob storage integration).
        """
        # Create temporary file
        with tempfile.NamedTemporaryFile(suffix=file_extension, delete=False) as temp_file:
            temp_file.write(file_bytes)
            temp_path = temp_file.name
        
        try:
            return self.extract_text(temp_path)
        finally:
            os.unlink(temp_path)
    
    # PDF Extraction Methods
    
    def _extract_pdf_pdfplumber(self, file_path: str) -> Dict[str, str]:
        """Extract PDF text using pdfplumber (best for complex layouts)."""
        text_parts = []
        formatted_parts = []
        
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                # Extract text with layout preservation
                page_text = page.extract_text(layout=True, x_tolerance=2, y_tolerance=2)
                if page_text:
                    text_parts.append(page_text)
                
                # Try table extraction for structured data
                tables = page.extract_tables()
                for table in tables:
                    table_text = self._format_table_as_text(table)
                    formatted_parts.append(table_text)
        
        raw_text = '\n'.join(text_parts + formatted_parts)
        cleaned_text = self._clean_extracted_text(raw_text)
        
        return {
            'raw_text': cleaned_text,
            'formatted_text': raw_text,
            'extraction_method': 'pdfplumber'
        }
    
    def _extract_pdf_pymupdf(self, file_path: str) -> Dict[str, str]:
        """Extract PDF text using PyMuPDF (good for text-heavy documents)."""
        doc = fitz.open(file_path)
        text_parts = []
        
        for page in doc:
            # Get text with better formatting
            text = page.get_text("text")
            if text.strip():
                text_parts.append(text)
        
        doc.close()
        
        raw_text = '\n'.join(text_parts)
        cleaned_text = self._clean_extracted_text(raw_text)
        
        return {
            'raw_text': cleaned_text,
            'formatted_text': raw_text,
            'extraction_method': 'pymupdf'
        }
    
    def _extract_pdf_pypdf2(self, file_path: str) -> Dict[str, str]:
        """Fallback PDF extraction using PyPDF2."""
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text_parts = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text.strip():
                    text_parts.append(text)
        
        raw_text = '\n'.join(text_parts)
        cleaned_text = self._clean_extracted_text(raw_text)
        
        return {
            'raw_text': cleaned_text,
            'formatted_text': raw_text,
            'extraction_method': 'pypdf2'
        }
    
    # DOCX Extraction Methods
    
    def _extract_docx_python_docx(self, file_path: str) -> Dict[str, str]:
        """Extract DOCX using python-docx library (preserves more structure)."""
        doc = Document(file_path)
        
        text_parts = []
        formatted_parts = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract tables
        for table in doc.tables:
            table_text = self._extract_table_from_docx(table)
            if table_text:
                formatted_parts.append(table_text)
        
        raw_text = '\n'.join(text_parts + formatted_parts)
        cleaned_text = self._clean_extracted_text(raw_text)
        
        return {
            'raw_text': cleaned_text,
            'formatted_text': raw_text,
            'extraction_method': 'python_docx'
        }
    
    def _extract_docx_docx2txt(self, file_path: str) -> Dict[str, str]:
        """Fallback DOCX extraction using docx2txt."""
        raw_text = docx2txt.process(file_path)
        cleaned_text = self._clean_extracted_text(raw_text)
        
        return {
            'raw_text': cleaned_text,
            'formatted_text': raw_text,
            'extraction_method': 'docx2txt'
        }
    
    # TXT Extraction
    
    def _extract_txt(self, file_path: str) -> Dict[str, str]:
        """Extract text from TXT files with encoding detection."""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    raw_text = file.read()
                cleaned_text = self._clean_extracted_text(raw_text)
                
                return {
                    'raw_text': cleaned_text,
                    'formatted_text': raw_text,
                    'extraction_method': f'txt_{encoding}'
                }
            except UnicodeDecodeError:
                continue
        
        raise Exception("Could not decode text file with any supported encoding")
    
    # Helper Methods
    
    def _clean_extracted_text(self, text: str) -> str:
        """
        Clean extracted text to fix bisection and formatting issues.
        """
        if not text:
            return ""
        
        # Fix common text bisection issues
        text = self._fix_text_bisection(text)
        
        # Clean up excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)  # Max 2 consecutive newlines
        text = re.sub(r' +', ' ', text)  # Multiple spaces to single space
        text = re.sub(r'\t+', ' ', text)  # Tabs to spaces
        
        # Fix broken words at line endings
        text = self._fix_line_break_words(text)
        
        # Normalize unicode characters
        text = text.encode('utf-8').decode('utf-8')
        
        return text.strip()
    
    def _fix_text_bisection(self, text: str) -> str:
        """
        Fix common text bisection patterns found in PDF extraction.
        """
        # Common bisection patterns
        bisection_patterns = [
            # Space in middle of words
            (r'([a-z])(\s+)([a-z])', r'\1\3'),
            # Broken compound words
            (r'([a-z])\s+([a-z]{1,2})\s+([a-z])', r'\1\2\3'),
            # Technology terms
            (r'Ja\s*v\s*a\s*Script', 'JavaScript'),
            (r'sen\s*ti\s*men\s*t', 'sentiment'),
            (r'mo\s*v\s*emen\s*t', 'movement'),
            (r'ey\s*e\s*s', 'eyes'),
            (r'universit\s*y', 'university'),
            (r'T\s*ec\s*hnology', 'Technology'),
            (r'programmin\s*g', 'programming'),
            # Email patterns
            (r'([a-zA-Z0-9]+)(\s+)@(\s+)([a-zA-Z0-9]+)', r'\1@\4'),
        ]
        
        for pattern, replacement in bisection_patterns:
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        
        return text
    
    def _fix_line_break_words(self, text: str) -> str:
        """
        Fix words broken across lines (hyphenation issues).
        """
        # Fix words broken with hyphens at line end
        text = re.sub(r'([a-z])-\s*\n\s*([a-z])', r'\1\2', text, flags=re.IGNORECASE)
        
        # Fix words broken without hyphens (common in PDF extraction)
        lines = text.split('\n')
        fixed_lines = []
        
        for i, line in enumerate(lines):
            if i < len(lines) - 1:
                next_line = lines[i + 1]
                
                # Check if current line ends with partial word and next line starts with continuation
                if (line and line[-1].islower() and 
                    next_line and next_line[0].islower() and
                    not line.endswith('.') and not line.endswith(',') and
                    len(line.split()[-1]) < 4):  # Likely broken word
                    
                    # Merge the lines
                    fixed_lines.append(line + next_line)
                    lines[i + 1] = ""  # Clear next line
                else:
                    fixed_lines.append(line)
            else:
                if line:  # Don't add empty last line
                    fixed_lines.append(line)
        
        return '\n'.join(line for line in fixed_lines if line)
    
    def _format_table_as_text(self, table: List[List[str]]) -> str:
        """Convert table data to readable text format."""
        if not table:
            return ""
        
        formatted_rows = []
        for row in table:
            if row and any(cell for cell in row if cell):  # Non-empty row
                clean_row = [cell.strip() if cell else "" for cell in row]
                formatted_rows.append(" | ".join(clean_row))
        
        return "\n".join(formatted_rows)
    
    def _extract_table_from_docx(self, table) -> str:
        """Extract table content from DOCX table object."""
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    cells.append(cell_text)
            if cells:
                rows.append(" | ".join(cells))
        
        return "\n".join(rows)
    
    def _assess_text_quality(self, text: str) -> float:
        """
        Assess the quality of extracted text (0-1 scale).
        Higher scores indicate better extraction quality.
        """
        if not text or len(text) < 50:
            return 0.0
        
        score = 1.0
        
        # Penalize excessive whitespace
        whitespace_ratio = len(re.findall(r'\s+', text)) / len(text)
        if whitespace_ratio > 0.3:
            score -= 0.2
        
        # Penalize broken words (single letters with spaces)
        broken_words = len(re.findall(r'\b[a-zA-Z]\s+[a-zA-Z]\s*', text))
        if broken_words > len(text.split()) * 0.1:
            score -= 0.3
        
        # Reward complete sentences
        sentences = re.findall(r'[.!?]+', text)
        if len(sentences) > 5:
            score += 0.1
        
        # Reward email patterns
        emails = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
        if emails:
            score += 0.1
        
        # Reward phone patterns
        phones = re.findall(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', text)
        if phones:
            score += 0.1
        
        return min(1.0, max(0.0, score))


# Global instance for easy import
enhanced_extractor = EnhancedTextExtractor()